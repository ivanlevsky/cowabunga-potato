from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.shared import qn
from docx.shared import Inches

import re

document = Document()


def add_document_text(text, **option):
    paragraph = document.add_paragraph('')
    default_font_name = 'Arial'
    default_font_size = 15
    default_font_bold = False
    default_font_italic = False
    if option.get('font_size') is not None:
        default_font_size = option.get('font_size')
    if option.get('font_bold') is not None:
        default_font_bold = option.get('font_bold')
    if option.get('font_italic') is not None:
        default_font_italic = option.get('font_italic')
    if option.get('font_name') is not None:
        default_font_name = option.get('font_name')
    runner = paragraph.add_run(text)
    runner.bold = default_font_bold
    runner.font.name = default_font_name
    runner.italic = default_font_italic
    runner.font.size = Pt(default_font_size)
    runner.add_break(WD_BREAK.LINE)
    return runner

def add_document_text_simple(text, **option):
    paragraph = document.add_paragraph('')
    default_font_size = 15
    if option.get('font_size') is not None:
        default_font_size = option.get('font_size')
    runner = paragraph.add_run(text)
    runner.font.size = Pt(default_font_size)
    runner.add_break(WD_BREAK.LINE)


def add_doc_table(table_datas, *styles):
    default_table_style = 'Table Grid'
    if styles.__len__() > 0:
        default_table_style = styles[0]
    col_num = 0
    row_num = table_datas.__len__()
    if row_num > 0:
        col_num = table_datas[0].__len__()
    table = document.add_table(rows=0, cols=col_num, style=default_table_style)
    for rt in table_datas:
        row_cells = table.add_row().cells
        for i in range(col_num):
            if rt[i] is None:
                row_cells[i].text = ''
            else:
                row_cells[i].text = str(rt[i])


def add_doc_picture(picture):
    # document.add_picture(picture, width=Inches(1.25))
    document.add_picture(picture)


def list_style():
    styles = [
        # s for s in document.styles if s.type == WD_STYLE_TYPE.CHARACTER
        # s for s in document.styles if s.type == WD_STYLE_TYPE.LIST
        # s for s in document.styles if s.type == WD_STYLE_TYPE.PARAGRAPH
        s for s in document.styles if s.type == WD_STYLE_TYPE.TABLE
    ]
    return styles


# convert text and keywords to list in order
def reformat_text_and_keyword(text):
    re_text = []
    split_text = list(filter(None, re.split('(\{[A-Za-z0-9_-]*\([A-Za-z0-9_-]*\)\})', text)))
    for st in split_text:
        keyword_and_params = list(filter(None, re.findall('\{[A-Za-z0-9_-]*\([A-Za-z0-9_-]*\)\}', st)))
        if keyword_and_params.__len__() > 0:
            for kp in keyword_and_params:
                kpl = list(filter(None, re.split('[{()}]', kp)))
                if kpl.__len__() > 0:
                    re_text.append(''.join((kpl[0],'_s_p_l_i_t_',kpl[1])))
                else:
                    re_text.append(st)
        else:
            # spl = list(filter(None, re.split('([ 	])', st)))
            # for sp in spl:
            re_text.append(st)
    return re_text


# read doc template replace text data and trigger keyword functions
def read_docx_template_file(file_name):
    read_document = Document(file_name)
    for doc_para in read_document.paragraphs:
        paragraph = document.add_paragraph('', style=doc_para.style)
        for rdt in reformat_text_and_keyword(doc_para.text):
            if rdt.__contains__('_s_p_l_i_t_'):
                if rdt.split('_s_p_l_i_t')[0] == 'time':
                    rdt = '2000-12-01'
                    paragraph.add_run(rdt)
                elif rdt.split('_s_p_l_i_t')[0] == 'table':
                    rdt = 'table_data1'
                    records = (
                        (3, 'Jack', None),
                        (7, 'Tom', '12'),
                        (4, 'Cloud', '33'),
                        (5, 'Zack', '1')
                    )
                    records = (('N0.', 'NAME', 'AGE'),) + records
                    add_doc_table(records, 'Light Grid Accent 5')
                else:
                    rdt = rdt
                    paragraph.add_run(rdt)
            else:
                for run in doc_para.runs:
                    runner = paragraph.add_run(run.text)
                    runner.font.name = run.font.name
                    runner.bold = run.bold
                    runner.italic = run.italic
                    runner.font.size = run.font.size
                    runner.underline = run.underline
                    runner.font.color.rgb = run.font.color.rgb
                # rdt = rdt
                # pr = paragraph.add_run(rdt)
        # paragraph.add_run(rdt)
        # paragraph.add_run().add_break(WD_BREAK.LINE)
        print('----end of para----')


def save_docx_file(file_name):
    document.save(file_name)


# generate all built-in table styles example
def write_docx_table_style_example():
    records = (
        (3, 'Jack', None),
        (7, 'Tom', '12'),
        (4, 'Cloud', '33'),
        (5, 'Zack', '1')
    )
    records = (('N0.', 'NAME', 'AGE'),) + records
    for style in list_style():
        add_document_text(style.name, font_size = 10, font_bold= False, font_italic = True)
        add_doc_table(records, style.name)



def write_docx_template_example():
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    # document.styles['Normal'].font.size = Pt(10.5)
    records = (
        ('序号', '名字', '年龄', '金额(元)'),
        (3, 'Jack', None, None),
        (7, 'Tom', '12', '1003.4'),
        (4, 'Cloud', '33', '338383.12'),
        (5, 'Zack', '12', '-383841')
    )

    add_document_text_simple("      段落1：测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱"
                             ,font_size=12)
    add_document_text_simple("      段落2：测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱"
                             ,font_size=12)
    add_doc_table(records, 'Light Grid Accent 5')
    add_document_text_simple("      段落3：测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱测试文字段乱"
                             ,font_size=12)






